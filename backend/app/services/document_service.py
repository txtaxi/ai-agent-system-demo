from io import BytesIO
from pathlib import Path
from uuid import uuid4

from docx import Document as DocxDocument
from fastapi import UploadFile
from pypdf import PdfReader
from sqlalchemy import delete, select
from sqlalchemy.orm import Session

from app.core.config import settings
from app.models.document import Document
from app.models.document_chunk import DocumentChunk
from app.models.knowledge_base import KnowledgeBase
from app.rag.chunking import chunk_document
from app.schemas.document import DocumentProcessResult
from app.services.embedding_service import generate_embedding


def _ensure_upload_dir() -> Path:
    upload_dir = Path(settings.upload_dir)
    upload_dir.mkdir(parents=True, exist_ok=True)
    return upload_dir


def save_uploaded_document(
    db: Session,
    knowledge_base_id: int,
    file: UploadFile,
) -> Document:
    # 上传阶段先把文件落盘并记录元数据，解析和分块放到后续处理步骤。
    knowledge_base = db.get(KnowledgeBase, knowledge_base_id)
    if knowledge_base is None:
        raise ValueError("知识库不存在。")

    upload_dir = _ensure_upload_dir()
    suffix = Path(file.filename or "").suffix
    stored_name = f"{uuid4().hex}{suffix}"
    file_path = upload_dir / stored_name

    with file_path.open("wb") as buffer:
        buffer.write(file.file.read())

    document = Document(
        knowledge_base_id=knowledge_base_id,
        filename=file.filename or stored_name,
        file_type=suffix.lstrip(".").lower() or "unknown",
        storage_path=str(file_path),
        status="uploaded",
    )
    db.add(document)
    db.commit()
    db.refresh(document)
    return document


def list_documents(db: Session, knowledge_base_id: int | None = None) -> list[Document]:
    stmt = select(Document).order_by(Document.created_at.desc())
    if knowledge_base_id is not None:
        stmt = stmt.where(Document.knowledge_base_id == knowledge_base_id)
    return list(db.scalars(stmt))


def list_document_chunks(db: Session, document_id: int) -> list[DocumentChunk]:
    document = db.get(Document, document_id)
    if document is None:
        raise ValueError("文档不存在。")

    stmt = (
        select(DocumentChunk)
        .where(DocumentChunk.document_id == document_id)
        .order_by(DocumentChunk.chunk_index.asc())
    )
    return list(db.scalars(stmt))


def _read_document_text(document: Document) -> str:
    file_path = Path(document.storage_path)
    if not file_path.exists():
        raise ValueError("文档文件不存在。")

    if document.file_type in {"txt", "md"}:
        return file_path.read_text(encoding="utf-8")
    if document.file_type == "pdf":
        return _read_pdf_text(file_path)
    if document.file_type == "docx":
        return _read_docx_text(file_path)

    raise ValueError("当前仅支持 txt、md、pdf 和 docx 文件处理。")


def _read_pdf_text(file_path: Path) -> str:
    reader = PdfReader(str(file_path))
    texts = [page.extract_text() or "" for page in reader.pages]
    content = "\n".join(texts).strip()
    if not content:
        raise ValueError("PDF 未提取到有效文本。")
    return content


def _read_docx_text(file_path: Path) -> str:
    with file_path.open("rb") as file:
        document = DocxDocument(BytesIO(file.read()))
    texts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    content = "\n".join(texts).strip()
    if not content:
        raise ValueError("DOCX 未提取到有效文本。")
    return content


def process_document(db: Session, document_id: int) -> DocumentProcessResult:
    # 文档处理主链路：读取文本 -> 分块 -> 向量化 -> 写入 document_chunks。
    document = db.get(Document, document_id)
    if document is None:
        raise ValueError("文档不存在。")

    document.status = "processing"
    document.error_message = None
    db.add(document)
    db.commit()
    db.refresh(document)

    try:
        content = _read_document_text(document)
        chunk_results = chunk_document(content)
        if not chunk_results:
            raise ValueError("文档没有可用于切分的有效文本。")

        db.execute(delete(DocumentChunk).where(DocumentChunk.document_id == document.id))

        for index, chunk in enumerate(chunk_results):
            metadata = {
                "source": document.filename,
                "file_type": document.file_type,
                **chunk.metadata,
            }
            db.add(
                DocumentChunk(
                    document_id=document.id,
                    chunk_index=index,
                    content=chunk.content,
                    metadata_json=metadata,
                    embedding=generate_embedding(chunk.content),
                )
            )

        document.status = "processed"
        document.error_message = None
        db.add(document)
        db.commit()
        db.refresh(document)

        return DocumentProcessResult(
            document_id=document.id,
            status=document.status,
            chunk_count=len(chunk_results),
            message="文档处理完成。",
        )
    except ValueError as exc:
        document.status = "failed"
        document.error_message = str(exc)
        db.add(document)
        db.commit()
        db.refresh(document)
        raise
    except UnicodeDecodeError as exc:
        document.status = "failed"
        document.error_message = "文件编码不受支持，请使用 UTF-8 编码。"
        db.add(document)
        db.commit()
        db.refresh(document)
        raise ValueError("文件编码不受支持，请使用 UTF-8 编码。") from exc
